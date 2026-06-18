import io

import docx
import pytest
from ai_employee.ingestion_worker.parsers import DocxParser


@pytest.fixture
def docx_bytes() -> bytes:
    """Generate a minimal valid DOCX in memory with headings and paragraphs."""
    doc = docx.Document()

    doc.add_heading("Network Fault Diagnosis", level=1)
    doc.add_paragraph("This is the introduction paragraph describing the fault.")

    doc.add_heading("Root Cause Analysis", level=2)
    doc.add_paragraph("The root cause was identified as a power failure.")
    doc.add_paragraph("Secondary cause: cooling system overload.")

    doc.add_heading("Resolution Steps", level=2)
    doc.add_paragraph("Step 1: Restart power module.")
    doc.add_paragraph("Step 2: Verify cooling system.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDocxParser:
    def test_builds_heading_hierarchy(self, docx_bytes: bytes) -> None:
        """Headings should produce section paths like 'H1 > H2'."""
        sections = DocxParser().parse(docx_bytes)
        paths = {s.section_path for s in sections}
        assert "Network Fault Diagnosis" in paths
        assert "Network Fault Diagnosis > Root Cause Analysis" in paths
        assert "Network Fault Diagnosis > Resolution Steps" in paths

    def test_paragraphs_become_blocks(self, docx_bytes: bytes) -> None:
        """Body paragraphs should appear as blocks in the correct section."""
        sections = DocxParser().parse(docx_bytes)
        all_blocks = [b for s in sections for b in s.blocks]
        combined = " ".join(all_blocks)
        assert "introduction paragraph" in combined
        assert "power failure" in combined
        assert "cooling system overload" in combined
        assert "Restart power module" in combined

    def test_blocks_under_correct_section(self, docx_bytes: bytes) -> None:
        """Paragraphs should be assigned to the nearest preceding heading."""
        sections = DocxParser().parse(docx_bytes)
        rca_sections = [s for s in sections if "Root Cause Analysis" in s.section_path]
        assert len(rca_sections) == 1
        rca_text = " ".join(rca_sections[0].blocks)
        assert "power failure" in rca_text

    def test_empty_docx_yields_empty(self) -> None:
        """An empty DOCX with no paragraphs should return empty list."""
        doc = docx.Document()
        buf = io.BytesIO()
        doc.save(buf)
        sections = DocxParser().parse(buf.getvalue())
        assert sections == []

    def test_docx_without_headings_uses_root(self, docx_bytes: bytes) -> None:
        """DOCX without any headings should use 'root' as section_path."""
        doc = docx.Document()
        doc.add_paragraph("Plain text without headings.")
        buf = io.BytesIO()
        doc.save(buf)
        sections = DocxParser().parse(buf.getvalue())
        assert len(sections) == 1
        assert sections[0].section_path == "root"
        assert sections[0].blocks == ["Plain text without headings."]
