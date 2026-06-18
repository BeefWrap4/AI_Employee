from __future__ import annotations

import io
from pathlib import Path

import pytest
from ai_employee.ingestion_worker.app import create_app
from fastapi.testclient import TestClient


def _write(tmp_path: Path, name: str, content: str) -> str:
    """Write text content as a file, return the path."""
    raw = tmp_path / "raw"
    raw.mkdir(exist_ok=True)
    path = raw / name
    path.write_text(content, encoding="utf-8")
    return str(path)


def _write_bytes(tmp_path: Path, name: str, content: bytes) -> str:
    """Write binary content as a file, return the path."""
    raw = tmp_path / "raw"
    raw.mkdir(exist_ok=True)
    path = raw / name
    path.write_bytes(content)
    return str(path)


def _set_data_dir(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    monkeypatch.setenv("KNOWLEDGE_DATA_DIR", str(tmp_path))


def test_health_reports_stub_provider() -> None:
    client = TestClient(create_app())
    resp = client.get("/health")
    assert resp.status_code == 200
    body = resp.json()
    assert body["service"] == "ingestion-worker"
    assert body["status"] == "ok"
    assert body["embedding_provider"] == "stub"


def test_parse_markdown_returns_chunks_and_embeddings(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    _set_data_dir(monkeypatch, tmp_path)
    file_path = _write(
        tmp_path,
        "doc_001.md",
        "# 接入排障\n## RRC\n先检查告警和 KPI。\n",
    )
    client = TestClient(create_app())
    resp = client.post(
        "/internal/parse",
        json={
            "doc_id": "doc_001",
            "file_path": file_path,
            "mime_type": "text/markdown",
            "metadata": {"network_type": "5g"},
        },
    )
    assert resp.status_code == 200
    body = resp.json()
    assert body["doc_id"] == "doc_001"
    assert body["embedding_model"] == "stub"
    assert len(body["chunks"]) >= 1
    assert len(body["embeddings"]) == len(body["chunks"])
    assert all(len(vec) == 8 for vec in body["embeddings"])
    assert body["chunks"][0]["chunk_id"].startswith("chunk_doc_001_")


def test_parse_unsupported_mime_returns_415(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    _set_data_dir(monkeypatch, tmp_path)
    file_path = _write(tmp_path, "doc.bin", "\x00\x01\x02")
    client = TestClient(create_app())
    resp = client.post(
        "/internal/parse",
        json={
            "doc_id": "doc_002",
            "file_path": file_path,
            "mime_type": "application/octet-stream",
            "metadata": {},
        },
    )
    assert resp.status_code == 415
    assert resp.json()["error_code"] == "mime_unsupported"


def test_parse_missing_file_returns_400(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    _set_data_dir(monkeypatch, tmp_path)
    (tmp_path / "raw").mkdir(exist_ok=True)
    client = TestClient(create_app())
    resp = client.post(
        "/internal/parse",
        json={
            "doc_id": "doc_003",
            "file_path": str(tmp_path / "raw" / "missing.md"),
            "mime_type": "text/markdown",
            "metadata": {},
        },
    )
    assert resp.status_code == 400
    assert resp.json()["detail"]["error_code"] == "file_not_found"


def test_parse_embedding_dim_consistent(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    _set_data_dir(monkeypatch, tmp_path)
    file_path = _write(tmp_path, "doc.md", "一段正文足够长以独立成块。")
    client = TestClient(create_app())
    resp = client.post(
        "/internal/parse",
        json={
            "doc_id": "doc_004",
            "file_path": file_path,
            "mime_type": "text/plain",
            "metadata": {},
        },
    )
    body = resp.json()
    dims = {len(v) for v in body["embeddings"]}
    assert len(dims) == 1


def test_parse_rejects_path_outside_data_dir(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    _set_data_dir(monkeypatch, tmp_path)
    (tmp_path / "raw").mkdir(exist_ok=True)
    safe = tmp_path / "raw" / "x.md"
    safe.write_text("safe", encoding="utf-8")

    outside = tmp_path / "outside.md"
    outside.write_text("outside", encoding="utf-8")

    client = TestClient(create_app())
    resp = client.post(
        "/internal/parse",
        json={
            "doc_id": "doc_safe",
            "file_path": str(safe),
            "mime_type": "text/plain",
            "metadata": {},
        },
    )
    assert resp.status_code == 200

    resp2 = client.post(
        "/internal/parse",
        json={
            "doc_id": "doc_outside",
            "file_path": str(outside),
            "mime_type": "text/plain",
            "metadata": {},
        },
    )
    assert resp2.status_code == 400
    assert resp2.json()["detail"]["error_code"] == "path_not_allowed"


def test_parse_rejects_relative_path(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    _set_data_dir(monkeypatch, tmp_path)
    client = TestClient(create_app())
    resp = client.post(
        "/internal/parse",
        json={
            "doc_id": "doc_rel",
            "file_path": "x.md",
            "mime_type": "text/plain",
            "metadata": {},
        },
    )
    assert resp.status_code == 400
    assert resp.json()["detail"]["error_code"] == "path_not_allowed"


# ---------------------------------------------------------------------------
# Binary parser integration tests (PDF / DOCX / XLSX)
# ---------------------------------------------------------------------------


class TestBinaryParsers:
    """End-to-end tests for PDF, DOCX, and XLSX parsers through the worker."""

    @staticmethod
    def _client(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> TestClient:
        _set_data_dir(monkeypatch, tmp_path)
        return TestClient(create_app())

    def test_parse_pdf_returns_chunks(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        import fitz

        buf = io.BytesIO()
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Network Fault Report", fontsize=18, fontname="helv")
        page.insert_text((72, 120), "Root cause: power failure.", fontsize=12, fontname="helv")
        doc.save(buf)
        doc.close()

        file_path = _write_bytes(tmp_path, "doc_001.pdf", buf.getvalue())
        client = self._client(tmp_path, monkeypatch)
        resp = client.post(
            "/internal/parse",
            json={
                "doc_id": "doc_001",
                "file_path": file_path,
                "mime_type": "application/pdf",
                "metadata": {},
            },
        )
        assert resp.status_code == 200
        body = resp.json()
        assert body["doc_id"] == "doc_001"
        assert len(body["chunks"]) >= 1
        assert any("Fault" in c["content"] for c in body["chunks"])

    def test_parse_docx_returns_chunks(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        import docx

        d = docx.Document()
        d.add_heading("Diagnosis", level=1)
        d.add_paragraph("The system experienced a fault.")
        buf = io.BytesIO()
        d.save(buf)

        file_path = _write_bytes(tmp_path, "doc_002.docx", buf.getvalue())
        client = self._client(tmp_path, monkeypatch)
        resp = client.post(
            "/internal/parse",
            json={
                "doc_id": "doc_002",
                "file_path": file_path,
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "metadata": {},
            },
        )
        assert resp.status_code == 200
        body = resp.json()
        assert len(body["chunks"]) >= 1
        assert any("fault" in c["content"].lower() for c in body["chunks"])

    def test_parse_xlsx_returns_chunks(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        import openpyxl

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Alarms"
        ws.append(["Time", "Event"])
        ws.append(["10:00", "Link Down"])
        buf = io.BytesIO()
        wb.save(buf)

        file_path = _write_bytes(tmp_path, "doc_003.xlsx", buf.getvalue())
        client = self._client(tmp_path, monkeypatch)
        resp = client.post(
            "/internal/parse",
            json={
                "doc_id": "doc_003",
                "file_path": file_path,
                "mime_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "metadata": {},
            },
        )
        assert resp.status_code == 200
        body = resp.json()
        assert len(body["chunks"]) >= 1
        assert any("Link Down" in c["content"] for c in body["chunks"])
