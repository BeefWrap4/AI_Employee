"""R33-C3: DOCX table parsing.

``DocxParser`` now iterates ``doc.tables`` in addition to
``doc.paragraphs``.  Each table produces a ``ParsedSection`` with
``section_path=f"Table {i}"``, ``table_id=f"docx_table_{i}"``, ``columns``
from the first row, ``values`` from subsequent rows, and text blocks
from joined cells (preserving the existing text-block representation).
"""

from __future__ import annotations

import io

import docx
from ai_employee.ingestion_worker.parsers import DocxParser


def _docx_with_table_bytes() -> bytes:
    doc = docx.Document()
    doc.add_paragraph("Intro paragraph before the table.")

    table = doc.add_table(rows=3, cols=3)
    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = "Time"
    hdr[1].text = "Event"
    hdr[2].text = "Severity"
    # Data rows
    r1 = table.rows[1].cells
    r1[0].text = "10:00"
    r1[1].text = "Link Down"
    r1[2].text = "Critical"
    r2 = table.rows[2].cells
    r2[0].text = "10:05"
    r2[1].text = "Failover"
    r2[2].text = "Warning"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDocxTableParser:
    def test_docx_with_table_yields_table_section(self) -> None:
        sections = DocxParser().parse(_docx_with_table_bytes())
        table_sections = [s for s in sections if s.section_path.startswith("Table ")]
        assert len(table_sections) == 1

    def test_docx_table_section_has_table_id(self) -> None:
        sections = DocxParser().parse(_docx_with_table_bytes())
        table_sections = [s for s in sections if s.section_path.startswith("Table ")]
        assert table_sections[0].table_id == "docx_table_0"

    def test_docx_table_section_has_columns(self) -> None:
        sections = DocxParser().parse(_docx_with_table_bytes())
        table_section = next(s for s in sections if s.section_path.startswith("Table "))
        assert table_section.columns == ["Time", "Event", "Severity"]

    def test_docx_table_section_has_values(self) -> None:
        sections = DocxParser().parse(_docx_with_table_bytes())
        table_section = next(s for s in sections if s.section_path.startswith("Table "))
        assert table_section.values is not None
        assert len(table_section.values) == 2
        assert table_section.values[0] == ["10:00", "Link Down", "Critical"]
        assert table_section.values[1] == ["10:05", "Failover", "Warning"]

    def test_docx_table_section_has_text_blocks(self) -> None:
        """Text blocks from joined cells are still emitted (backward compat)."""
        sections = DocxParser().parse(_docx_with_table_bytes())
        table_section = next(s for s in sections if s.section_path.startswith("Table "))
        assert len(table_section.blocks) == 2
        assert "Link Down" in table_section.blocks[0]
        assert "Failover" in table_section.blocks[1]

    def test_docx_paragraphs_still_parsed(self) -> None:
        """Existing paragraph parsing is not broken by table support."""
        sections = DocxParser().parse(_docx_with_table_bytes())
        all_blocks = [
            b for s in sections for b in s.blocks if not s.section_path.startswith("Table ")
        ]
        assert any("Intro paragraph" in b for b in all_blocks)

    def test_docx_multiple_tables_indexed(self) -> None:
        doc = docx.Document()
        t1 = doc.add_table(rows=2, cols=2)
        t1.rows[0].cells[0].text = "A"
        t1.rows[0].cells[1].text = "B"
        t1.rows[1].cells[0].text = "1"
        t1.rows[1].cells[1].text = "2"
        t2 = doc.add_table(rows=2, cols=2)
        t2.rows[0].cells[0].text = "C"
        t2.rows[0].cells[1].text = "D"
        t2.rows[1].cells[0].text = "3"
        t2.rows[1].cells[1].text = "4"
        buf = io.BytesIO()
        doc.save(buf)
        sections = DocxParser().parse(buf.getvalue())
        table_sections = [s for s in sections if s.section_path.startswith("Table ")]
        assert len(table_sections) == 2
        ids = {s.table_id for s in table_sections}
        assert ids == {"docx_table_0", "docx_table_1"}
        # section_path indexes match table_ids
        by_path = {s.section_path: s for s in table_sections}
        assert by_path["Table 0"].table_id == "docx_table_0"
        assert by_path["Table 1"].table_id == "docx_table_1"
        assert by_path["Table 1"].columns == ["C", "D"]

    def test_docx_without_tables_unchanged(self) -> None:
        """A DOCX with no tables still parses like before (no table sections)."""
        doc = docx.Document()
        doc.add_paragraph("Just a paragraph.")
        buf = io.BytesIO()
        doc.save(buf)
        sections = DocxParser().parse(buf.getvalue())
        assert all(not s.section_path.startswith("Table ") for s in sections)
        assert all(s.table_id is None for s in sections)
